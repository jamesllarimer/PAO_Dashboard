from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colour palette ────────────────────────────────────────────────────────────
ARMY_GREEN  = RGBColor(0x4B, 0x52, 0x20)   # dark olive
METHOD_COLORS = {
    "GET":    RGBColor(0x1A, 0x56, 0x76),   # blue
    "POST":   RGBColor(0x15, 0x6B, 0x37),   # green
    "PUT":    RGBColor(0x7A, 0x4F, 0x00),   # amber
    "DELETE": RGBColor(0x8B, 0x22, 0x15),   # red
}
HEADER_BG   = "2F3A1A"  # dark green for table headers (hex, no #)
ALT_ROW_BG  = "F5F5F0"  # very light grey for alternating rows


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_table_header_row(table, headers: list[str], col_widths: list[float]):
    row = table.rows[0]
    for i, (cell, header, width) in enumerate(zip(row.cells, headers, col_widths)):
        cell.width = Inches(width)
        set_cell_bg(cell, HEADER_BG)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)


def add_data_row(table, values: list[str], alt: bool = False):
    row = table.add_row()
    for i, (cell, value) in enumerate(zip(row.cells, values)):
        if alt:
            set_cell_bg(cell, ALT_ROW_BG)
        p = cell.paragraphs[0]
        run = p.add_run(value)
        run.font.size = Pt(9)
    return row


def method_badge(para, method: str):
    run = para.add_run(f" {method} ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # background shading on the run is not natively supported in python-docx;
    # we colour the text instead and note the method colour
    run.font.color.rgb = METHOD_COLORS.get(method, RGBColor(0x33, 0x33, 0x33))


def add_endpoint_heading(doc, method: str, path: str, description: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run_method = p.add_run(f"[{method}]  ")
    run_method.bold = True
    run_method.font.size = Pt(11)
    run_method.font.color.rgb = METHOD_COLORS.get(method, RGBColor(0x33, 0x33, 0x33))
    run_path = p.add_run(path)
    run_path.bold = True
    run_path.font.size = Pt(11)
    run_path.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    if description:
        p2 = doc.add_paragraph(description)
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after = Pt(4)
        p2.runs[0].font.size = Pt(9)
        p2.runs[0].italic = True


def add_fields_table(doc, title: str, rows: list[tuple]):
    """rows = list of (Field, Type, Required, Notes)  OR  (Field, Type, Notes)"""
    if not rows:
        return
    p = doc.add_paragraph(title)
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(9)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)

    four_col = len(rows[0]) == 4
    headers = ["Field", "Type", "Required", "Notes"] if four_col else ["Field", "Type", "Notes"]
    widths  = [1.4, 1.2, 0.9, 2.5] if four_col else [1.7, 1.3, 3.0]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_table_header_row(table, headers, widths)

    for idx, row_data in enumerate(rows):
        add_data_row(table, list(row_data), alt=(idx % 2 == 1))

    doc.add_paragraph()   # spacer


# ── Document data ─────────────────────────────────────────────────────────────

SUMMARY = [
    # (Method, Path, Description)
    ("GET",    "/api/v1/event_type",                  "Get all event types"),
    ("GET",    "/api/v1/event_type/{id}",             "Get event type by ID"),
    ("POST",   "/api/v1/event_type",                  "Create a new event type"),
    ("GET",    "/api/v1/event_status",                "Get all event statuses"),
    ("POST",   "/api/v1/event_status",                "Create a new event status"),
    ("GET",    "/api/v1/user",                        "Get all user profiles"),
    ("GET",    "/api/v1/user/{id}",                   "Get user profile by ID"),
    ("GET",    "/api/v1/product_type",                "Get all product types"),
    ("POST",   "/api/v1/product_type",                "Create a new product type"),
    ("GET",    "/api/v1/theme",                       "Get all themes"),
    ("GET",    "/api/v1/theme/{id}",                  "Get theme by ID"),
    ("POST",   "/api/v1/theme",                       "Create a new theme"),
    ("DELETE", "/api/v1/theme/{id}",                  "Delete a theme by ID"),
    ("GET",    "/api/v1/posting_locations",           "Get all posting locations"),
    ("POST",   "/api/v1/posting_locations",           "Create a new posting location"),
    ("GET",    "/api/v1/events",                      "Get all events"),
    ("GET",    "/api/v1/events/{id}",                 "Get event by ID"),
    ("GET",    "/api/v1/events/userId/{userId}",      "Get events by user ID"),
    ("POST",   "/api/v1/events",                      "Create a new event"),
    ("PUT",    "/api/v1/events/{id}",                 "Update an existing event"),
    ("DELETE", "/api/v1/events/{id}/delete",          "Delete an event by ID"),
    ("GET",    "/api/v1/theme_example",               "Get all theme examples"),
    ("POST",   "/api/v1/theme_example",               "Create a new theme example"),
]

# Field rows: (Field, Type, Required, Notes)
EVENT_REQUEST = [
    ("name",              "String",  "Yes", "Display name of the event"),
    ("description",       "String",  "Yes", "Event description / coverage plan"),
    ("startDate",         "String",  "Yes", "Format: yyyy-MM-dd"),
    ("endDate",           "String",  "Yes", "Format: yyyy-MM-dd"),
    ("leadId",            "Long",    "No",  "FK → UserProfile.id"),
    ("eventStatusId",     "Long",    "No",  "FK → EventStatus.id"),
    ("eventTypeId",       "Long",    "No",  "FK → EventType.id"),
    ("productTypeId",     "Long",    "No",  "FK → ProductType.id"),
    ("postingLocationId", "Long",    "No",  "FK → PostingLocation.id"),
    ("eventThemeId",      "Long",    "No",  "FK → Theme.id"),
]

EVENT_RESPONSE_DTO = [
    ("id",                 "Long",   "Auto-generated primary key"),
    ("name",               "String", ""),
    ("description",        "String", ""),
    ("eventType",          "String", "Resolved name of the event type"),
    ("eventTypeId",        "Long",   ""),
    ("startDate",          "String", "yyyy-MM-dd"),
    ("endDate",            "String", "yyyy-MM-dd"),
    ("lead",               "String", "Resolved full name of lead user"),
    ("leadId",             "Long",   ""),
    ("unit",               "String", "Resolved unit name of lead user"),
    ("unitId",             "Long",   ""),
    ("status",             "String", "Resolved event status name"),
    ("eventStatusId",      "Long",   ""),
    ("theme",              "String", "Resolved theme name"),
    ("eventThemeId",       "Long",   ""),
    ("postingLocation",    "String", "Resolved posting location name"),
    ("postingLocationId",  "Long",   ""),
    ("productType",        "String", "Resolved product type name"),
    ("productTypeId",      "Long",   ""),
]

EVENT_TYPE_FIELDS = [
    ("id",          "Long",   "Auto-generated"),
    ("name",        "String", ""),
    ("description", "String", ""),
]

EVENT_STATUS_FIELDS = [
    ("id",          "Long",   "Auto-generated"),
    ("name",        "String", ""),
    ("description", "String", ""),
]

USER_RESPONSE_FIELDS = [
    ("id",               "Long",   ""),
    ("username",         "String", ""),
    ("firstName",        "String", ""),
    ("lastName",         "String", ""),
    ("role",             "String", "PAO_UNIT or HQ_VIEWER"),
    ("unitName",         "String", ""),
    ("rankAbbreviation", "String", "e.g. SSG, CPT"),
]

PRODUCT_TYPE_FIELDS = [
    ("id",          "Long",   "Auto-generated"),
    ("name",        "String", ""),
    ("description", "String", ""),
]

THEME_REQUEST_FIELDS = [
    ("name",          "String",          "Yes", "Theme name"),
    ("theme_examples","List<ThemeExample>","No", "Embedded examples"),
]

THEME_EXAMPLE_FIELDS_NESTED = [
    ("name",        "String", ""),
    ("description", "String", ""),
]

THEME_RESPONSE_FIELDS = [
    ("id",             "Long",             "Auto-generated"),
    ("name",           "String",           ""),
    ("theme_examples", "List<ThemeExample>","Nested array; each item has id, name, description"),
]

POSTING_LOCATION_FIELDS = [
    ("id",   "Long",   "Auto-generated"),
    ("name", "String", ""),
]

THEME_EXAMPLE_REQUEST = [
    ("name",        "String", "Yes", ""),
    ("description", "String", "Yes", ""),
    ("theme",       "Object", "Yes", "Must include { \"id\": <themeId> }"),
]

THEME_EXAMPLE_RESPONSE = [
    ("id",          "Long",   "Auto-generated"),
    ("name",        "String", ""),
    ("description", "String", ""),
]


# ── Build document ────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # ── Title ──────────────────────────────────────────────────────────────────
    title = doc.add_heading("PAO Dashboard — REST API Reference", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = ARMY_GREEN

    sub = doc.add_paragraph("Base URL:  http://localhost:8080")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(10)
    sub.runs[0].italic = True

    doc.add_paragraph()

    # ── Summary table ──────────────────────────────────────────────────────────
    doc.add_heading("Endpoint Summary", level=1).runs[0].font.color.rgb = ARMY_GREEN

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_table_header_row(tbl, ["Method", "Path", "Description"], [0.8, 2.8, 3.0])

    for idx, (method, path, desc) in enumerate(SUMMARY):
        row = tbl.add_row()
        # Method cell
        set_cell_bg(row.cells[0], ALT_ROW_BG if idx % 2 else "FFFFFF")
        p_m = row.cells[0].paragraphs[0]
        p_m.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_m.add_run(method)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = METHOD_COLORS.get(method, RGBColor(0x33, 0x33, 0x33))
        # Path
        set_cell_bg(row.cells[1], ALT_ROW_BG if idx % 2 else "FFFFFF")
        rp = row.cells[1].paragraphs[0].add_run(path)
        rp.font.size = Pt(9)
        # Desc
        set_cell_bg(row.cells[2], ALT_ROW_BG if idx % 2 else "FFFFFF")
        rd = row.cells[2].paragraphs[0].add_run(desc)
        rd.font.size = Pt(9)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 1. EVENT TYPE
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("1. Event Type  —  /api/v1/event_type", level=1).runs[0].font.color.rgb = ARMY_GREEN

    add_endpoint_heading(doc, "GET", "/api/v1/event_type", "Returns a list of all event types.")
    add_fields_table(doc, "Response Body — List<EventType>:", EVENT_TYPE_FIELDS)

    add_endpoint_heading(doc, "GET", "/api/v1/event_type/{id}", "Returns a single event type by its ID.")
    p = doc.add_paragraph(); p.add_run("Path Variable: ").bold = True; p.add_run("id (Long) — EventType ID").font.size = Pt(9)
    add_fields_table(doc, "Response Body — EventType:", EVENT_TYPE_FIELDS)

    add_endpoint_heading(doc, "POST", "/api/v1/event_type", "Creates a new event type.")
    add_fields_table(doc, "Request Body — EventType:", [
        ("name",        "String", "Yes", ""),
        ("description", "String", "No",  ""),
    ])
    add_fields_table(doc, "Response Body — EventType:", EVENT_TYPE_FIELDS)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 2. EVENT STATUS
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("2. Event Status  —  /api/v1/event_status", level=1).runs[0].font.color.rgb = ARMY_GREEN

    add_endpoint_heading(doc, "GET", "/api/v1/event_status", "Returns a list of all event statuses.")
    add_fields_table(doc, "Response Body — List<EventStatus>:", EVENT_STATUS_FIELDS)

    add_endpoint_heading(doc, "POST", "/api/v1/event_status", "Creates a new event status.")
    add_fields_table(doc, "Request Body — EventStatus:", [
        ("name",        "String", "Yes", ""),
        ("description", "String", "No",  ""),
    ])
    add_fields_table(doc, "Response Body — EventStatus:", EVENT_STATUS_FIELDS)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 3. USER
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("3. User Profile  —  /api/v1/user", level=1).runs[0].font.color.rgb = ARMY_GREEN

    add_endpoint_heading(doc, "GET", "/api/v1/user", "Returns a list of all user profiles.")
    add_fields_table(doc, "Response Body — List<UserProfileResponseDto>:", USER_RESPONSE_FIELDS)

    add_endpoint_heading(doc, "GET", "/api/v1/user/{id}", "Returns a single user profile by ID.")
    p = doc.add_paragraph(); p.add_run("Path Variable: ").bold = True; p.add_run("id (Long) — UserProfile ID").font.size = Pt(9)
    add_fields_table(doc, "Response Body — UserProfileResponseDto:", USER_RESPONSE_FIELDS)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 4. PRODUCT TYPE
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("4. Product Type  —  /api/v1/product_type", level=1).runs[0].font.color.rgb = ARMY_GREEN

    add_endpoint_heading(doc, "GET", "/api/v1/product_type", "Returns a list of all product types.")
    add_fields_table(doc, "Response Body — List<ProductType>:", PRODUCT_TYPE_FIELDS)

    add_endpoint_heading(doc, "POST", "/api/v1/product_type", "Creates a new product type.")
    add_fields_table(doc, "Request Body — ProductType:", [
        ("name",        "String", "Yes", ""),
        ("description", "String", "No",  ""),
    ])
    add_fields_table(doc, "Response Body — ProductType:", PRODUCT_TYPE_FIELDS)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 5. THEME
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("5. Theme  —  /api/v1/theme", level=1).runs[0].font.color.rgb = ARMY_GREEN

    add_endpoint_heading(doc, "GET", "/api/v1/theme", "Returns all themes including their nested examples.")
    add_fields_table(doc, "Response Body — List<Theme>:", THEME_RESPONSE_FIELDS)

    add_endpoint_heading(doc, "GET", "/api/v1/theme/{id}", "Returns a single theme by ID.")
    p = doc.add_paragraph(); p.add_run("Path Variable: ").bold = True; p.add_run("id (Long) — Theme ID").font.size = Pt(9)
    add_fields_table(doc, "Response Body — Theme:", THEME_RESPONSE_FIELDS)

    add_endpoint_heading(doc, "POST", "/api/v1/theme", "Creates a new theme, optionally with nested examples.")
    add_fields_table(doc, "Request Body — Theme:", THEME_REQUEST_FIELDS)
    add_fields_table(doc, "Nested ThemeExample object:", [
        ("name",        "String", "Yes", ""),
        ("description", "String", "Yes", ""),
    ])
    add_fields_table(doc, "Response Body — Theme:", THEME_RESPONSE_FIELDS)

    add_endpoint_heading(doc, "DELETE", "/api/v1/theme/{id}", "Deletes a theme by ID. Returns 204 No Content. Will fail if the theme is assigned to an active event.")
    p = doc.add_paragraph(); p.add_run("Path Variable: ").bold = True; p.add_run("id (Long) — Theme ID").font.size = Pt(9)
    p2 = doc.add_paragraph("Response: 204 No Content (empty body)")
    p2.runs[0].italic = True; p2.runs[0].font.size = Pt(9)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 6. POSTING LOCATION
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("6. Posting Location  —  /api/v1/posting_locations", level=1).runs[0].font.color.rgb = ARMY_GREEN

    add_endpoint_heading(doc, "GET", "/api/v1/posting_locations", "Returns a list of all posting locations.")
    add_fields_table(doc, "Response Body — List<PostingLocation>:", POSTING_LOCATION_FIELDS)

    add_endpoint_heading(doc, "POST", "/api/v1/posting_locations", "Creates a new posting location.")
    add_fields_table(doc, "Request Body — PostingLocation:", [
        ("name", "String", "Yes", ""),
    ])
    add_fields_table(doc, "Response Body — PostingLocation:", POSTING_LOCATION_FIELDS)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 7. EVENT
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("7. Event  —  /api/v1/events", level=1).runs[0].font.color.rgb = ARMY_GREEN

    add_endpoint_heading(doc, "GET", "/api/v1/events", "Returns all events across all users as EventResponseDto objects.")
    add_fields_table(doc, "Response Body — List<EventResponseDto>:", EVENT_RESPONSE_DTO)

    add_endpoint_heading(doc, "GET", "/api/v1/events/{id}", "Returns a single event by ID.")
    p = doc.add_paragraph(); p.add_run("Path Variable: ").bold = True; p.add_run("id (Long) — Event ID").font.size = Pt(9)
    add_fields_table(doc, "Response Body — Event (entity with resolved FK objects):", [
        ("id",             "Long",          ""),
        ("name",           "String",        ""),
        ("description",    "String",        ""),
        ("event_type",     "EventType",     "Nested object"),
        ("start_date",     "String",        "yyyy-MM-dd"),
        ("end_date",       "String",        "yyyy-MM-dd"),
        ("lead",           "UserProfile",   "Nested object"),
        ("eventStatus",    "EventStatus",   "Nested object"),
        ("theme",          "Theme",         "Nested object"),
        ("postingLocation","PostingLocation","Nested object"),
        ("productType",    "ProductType",   "Nested object"),
    ])

    add_endpoint_heading(doc, "GET", "/api/v1/events/userId/{userId}", "Returns all events belonging to the given user.")
    p = doc.add_paragraph(); p.add_run("Path Variable: ").bold = True; p.add_run("userId (Long) — UserProfile ID").font.size = Pt(9)
    add_fields_table(doc, "Response Body — List<EventResponseDto>:", EVENT_RESPONSE_DTO)

    add_endpoint_heading(doc, "POST", "/api/v1/events", "Creates a new event. All FK fields accept the numeric ID of the related entity.")
    add_fields_table(doc, "Request Body — EventRequest:", EVENT_REQUEST)
    add_fields_table(doc, "Response Body — Event (entity):", [
        ("id",             "Long",          "Auto-generated"),
        ("name",           "String",        ""),
        ("description",    "String",        ""),
        ("event_type",     "EventType",     "Nested object"),
        ("start_date",     "String",        "yyyy-MM-dd"),
        ("end_date",       "String",        "yyyy-MM-dd"),
        ("lead",           "UserProfile",   "Nested object"),
        ("eventStatus",    "EventStatus",   "Nested object"),
        ("theme",          "Theme",         "Nested object"),
        ("postingLocation","PostingLocation","Nested object"),
        ("productType",    "ProductType",   "Nested object"),
    ])

    add_endpoint_heading(doc, "PUT", "/api/v1/events/{id}", "Updates an existing event. Accepts the same body as POST.")
    p = doc.add_paragraph(); p.add_run("Path Variable: ").bold = True; p.add_run("id (Long) — Event ID").font.size = Pt(9)
    add_fields_table(doc, "Request Body — EventRequest:", EVENT_REQUEST)
    add_fields_table(doc, "Response Body — EventResponseDto:", EVENT_RESPONSE_DTO)

    add_endpoint_heading(doc, "DELETE", "/api/v1/events/{id}/delete", "Deletes an event by ID. Returns 204 No Content.")
    p = doc.add_paragraph(); p.add_run("Path Variable: ").bold = True; p.add_run("id (Long) — Event ID").font.size = Pt(9)
    p2 = doc.add_paragraph("Response: 204 No Content (empty body)")
    p2.runs[0].italic = True; p2.runs[0].font.size = Pt(9)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 8. THEME EXAMPLE
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("8. Theme Example  —  /api/v1/theme_example", level=1).runs[0].font.color.rgb = ARMY_GREEN

    add_endpoint_heading(doc, "GET", "/api/v1/theme_example", "Returns all theme examples.")
    add_fields_table(doc, "Response Body — List<ThemeExample>:", THEME_EXAMPLE_RESPONSE)

    add_endpoint_heading(doc, "POST", "/api/v1/theme_example", "Creates a standalone theme example associated with an existing theme.")
    add_fields_table(doc, "Request Body — ThemeExample:", THEME_EXAMPLE_REQUEST)
    add_fields_table(doc, "Response Body — ThemeExample:", THEME_EXAMPLE_RESPONSE)

    # ── Save ──────────────────────────────────────────────────────────────────
    out = "api_docs.docx"
    doc.save(out)
    print(f"Saved → {out}")


if __name__ == "__main__":
    build()
